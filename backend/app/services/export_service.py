"""
导出服务 / Export service

为章节正文提供轻量导出能力，首版支持 TXT、Markdown、DOCX。
Provides lightweight chapter export for TXT, Markdown, and DOCX formats.
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
from typing import Dict, List, Optional, Sequence

import yaml
from docx import Document

from app.dependencies import get_draft_storage


@dataclass
class ExportChapter:
    """导出章节数据 / Export chapter payload."""

    chapter_id: str
    title: str
    content: str
    volume_id: Optional[str] = None


@dataclass
class ExportArtifact:
    """导出结果 / Export artifact."""

    filename: str
    media_type: str
    content: bytes


class ExportService:
    """章节导出服务 / Chapter export service."""

    MEDIA_TYPES: Dict[str, str] = {
        "txt": "text/plain; charset=utf-8",
        "md": "text/markdown; charset=utf-8",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    def __init__(self):
        self.draft_storage = get_draft_storage()

    async def export_chapters(
        self,
        project_id: str,
        chapter_ids: Sequence[str],
        fmt: str,
        include_chapter_titles: bool = True,
    ) -> ExportArtifact:
        """导出指定章节 / Export selected chapters."""
        normalized_format = str(fmt or "").strip().lower()
        if normalized_format not in self.MEDIA_TYPES:
            raise ValueError(f"Unsupported export format: {fmt}")

        chapters = await self._collect_chapters(project_id, chapter_ids)
        if not chapters:
            raise ValueError("No exportable chapters found")

        project_name = self._read_project_name(project_id)
        filename = self._build_filename(project_name, chapters, normalized_format)
        renderer = getattr(self, f"_render_{normalized_format}")
        content = renderer(chapters, include_chapter_titles=include_chapter_titles)
        return ExportArtifact(
            filename=filename,
            media_type=self.MEDIA_TYPES[normalized_format],
            content=content,
        )

    async def _collect_chapters(self, project_id: str, chapter_ids: Sequence[str]) -> List[ExportChapter]:
        """按现有项目顺序收集章节正文 / Collect chapters in persisted project order."""
        requested = [str(ch or "").strip() for ch in (chapter_ids or []) if str(ch or "").strip()]
        all_chapters = await self.draft_storage.list_chapters(project_id)
        ordered_ids = [chapter for chapter in all_chapters if not requested or chapter in set(requested)]

        summaries = await self.draft_storage.list_chapter_summaries(project_id)
        summary_map = {str(item.chapter): item for item in summaries}

        chapters: List[ExportChapter] = []
        for chapter_id in ordered_ids:
            content = await self.draft_storage.get_final_draft(project_id, chapter_id)
            if content is None:
                latest = await self.draft_storage.get_latest_draft(project_id, chapter_id)
                content = latest.content if latest else ""

            summary = summary_map.get(str(chapter_id))
            title = ""
            volume_id = None
            if summary:
                title = str(summary.title or "").strip()
                volume_id = getattr(summary, "volume_id", None)

            chapters.append(
                ExportChapter(
                    chapter_id=str(chapter_id),
                    title=title or str(chapter_id),
                    content=str(content or ""),
                    volume_id=volume_id,
                )
            )
        return chapters

    def _render_txt(self, chapters: Sequence[ExportChapter], include_chapter_titles: bool = True) -> bytes:
        """渲染 TXT / Render TXT."""
        parts: List[str] = []
        for index, chapter in enumerate(chapters):
            if index > 0:
                parts.append("")
                parts.append("")
            if include_chapter_titles:
                parts.append(chapter.title)
                parts.append("")
            parts.append(chapter.content.rstrip())
        return "\n".join(parts).strip().encode("utf-8")

    def _render_md(self, chapters: Sequence[ExportChapter], include_chapter_titles: bool = True) -> bytes:
        """渲染 Markdown / Render Markdown."""
        parts: List[str] = []
        for index, chapter in enumerate(chapters):
            if index > 0:
                parts.append("")
                parts.append("---")
                parts.append("")
            if include_chapter_titles:
                parts.append(f"## {chapter.title}")
                parts.append("")
            parts.append(chapter.content.rstrip())
        return "\n".join(parts).strip().encode("utf-8")

    def _render_docx(self, chapters: Sequence[ExportChapter], include_chapter_titles: bool = True) -> bytes:
        """渲染 DOCX / Render DOCX."""
        document = Document()
        for index, chapter in enumerate(chapters):
            if index > 0:
                document.add_page_break()
            if include_chapter_titles:
                document.add_heading(chapter.title, level=1)
            self._append_docx_paragraphs(document, chapter.content)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _append_docx_paragraphs(document: Document, content: str) -> None:
        """按段落写入 DOCX / Append DOCX paragraphs with simple blank-line splitting."""
        text = str(content or "").replace("\r\n", "\n")
        paragraphs = text.split("\n")
        for paragraph in paragraphs:
            if paragraph.strip():
                document.add_paragraph(paragraph)
            else:
                document.add_paragraph("")

    def _read_project_name(self, project_id: str) -> str:
        """读取项目名 / Read project name for exported filename."""
        project_path = self.draft_storage.get_project_path(project_id) / "project.yaml"
        if not project_path.exists():
            return str(project_id)
        try:
            data = yaml.safe_load(project_path.read_text(encoding="utf-8")) or {}
            return str(data.get("name") or project_id)
        except Exception:
            return str(project_id)

    def _build_filename(self, project_name: str, chapters: Sequence[ExportChapter], fmt: str) -> str:
        """构建导出文件名 / Build export filename."""
        base = self._safe_filename(project_name) or "wenshape_export"
        if len(chapters) == 1:
            chapter_name = self._safe_filename(chapters[0].title or chapters[0].chapter_id)
            if chapter_name:
                return f"{base}_{chapter_name}.{fmt}"
        return f"{base}_export.{fmt}"

    @staticmethod
    def _safe_filename(name: str) -> str:
        """清洗文件名 / Sanitize filename for Windows-safe output."""
        cleaned = re.sub(r'[<>:"/\\\\|?*]+', "_", str(name or "").strip())
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
        return cleaned[:120]


export_service = ExportService()
